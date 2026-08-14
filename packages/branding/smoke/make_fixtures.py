#!/usr/bin/env python3
"""[fork-only] 生成 UI 测试专用项目与各格式样本文件 [feat: ui-probe-toolkit] 2026-08-13

## 为什么要有独立测试项目

CHECKLIST 第 3 组含**归档 / 删除 / 分享**这类破坏性或对外的动作,第 4 组要逐格式验预览。
在 user 的真实项目(如 /Volumes/ExtSSD/Finance)里跑这些 = 拿真实数据练手,
「分享」还会把内容发到站外。所以固定用一个自建项目,内容全是可丢弃的样本。

顺带解决第 4 组的取材问题:.docx/.xlsx/.pdf/图片/超大文件 每次手工找一遍很浪费,
这里一次性生成,且**内容带可判定特征**(如 md 里的粗体/引用/内链),
预览对不对能直接断言,而不是「看着像渲染出来了」。

生成的是 git 仓且**故意留有未提交改动** —— 第 2 组 #15 的「N 更改」tab 需要真有 diff 才渲染。

跑法:
    python3 packages/branding/smoke/make_fixtures.py [目标目录]
默认目标 /Volumes/ExtSSD/deskfox-uitest(与 user 真实项目分开)。
"""
import os
import subprocess
import sys

# FORK 2026-08-14 [feat: upstream-sync-2026-08]:两处路径按平台解析,不再写死 mac。
# 原实现把测试项目目录和 soffice 路径都硬编码成 macOS 的,在 Windows 上跑会
# ① 把项目建到一个不存在的 /Volumes 路径下,② PDF 生成静默跳过 —— 而第 4 组预览用例
# 恰恰要用那个 PDF,于是「没有样本」会一路顺延成「预览验不了」。
HERE = os.path.dirname(os.path.abspath(__file__))
IS_WIN = sys.platform == "win32"

_DEFAULT_ROOT = "D:\\deskfox-uitest" if IS_WIN else "/Volumes/ExtSSD/deskfox-uitest"
ROOT = sys.argv[1] if len(sys.argv) > 1 else _DEFAULT_ROOT


def _find_soffice() -> str:
    """找一个能用的 soffice。

    优先用**打包产物内置的那份**(和用户实际用到的是同一份,能顺带验证 bundle 健康);
    找不到再回落到系统安装的 LibreOffice。两者都没有时返回 ""(生成 PDF 那步会明说跳过原因)。
    """
    unpacked = os.path.abspath(os.path.join(HERE, "..", "..", "desktop", "dist-deskfox"))
    candidates = []
    if IS_WIN:
        candidates += [
            os.path.join(unpacked, "win-unpacked", "libreoffice", "program", "soffice.exe"),
            os.path.abspath(os.path.join(HERE, "..", "libreoffice-bundle", "windows", "program", "soffice.exe")),
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        ]
    else:
        candidates += [
            os.path.join(unpacked, "mac-arm64", "DeskFox 本地版.app", "Contents", "Resources",
                         "libreoffice", "Contents", "MacOS", "soffice"),
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]
    for c in candidates:
        if os.path.exists(c):
            return c
    return ""


APP = _find_soffice()

# 每个样本都带**可判定特征**,预览结果能断言而不是靠肉眼「看着对」
MD_MAIN = """# 预览验收样本

这是一段普通正文,含**加粗特征词 BOLDMARK** 与 `行内代码 CODEMARK`。

## 二级标题 HEAD2

> 引用块特征词 QUOTEMARK

- 列表项一 LIST1
- 列表项二 LIST2
- 列表项三 LIST3

1. 有序项一
2. 有序项二

站内链接:[跳到子文档](./notes/sub.md)
站外链接:[example](https://example.com)

| 列 A | 列 B |
|---|---|
| A1 | B1 |
| A2 | B2 |
"""

MD_SUB = """# 子文档 SUBDOC

从 README 的站内链接跳过来应当**留在应用内**,不得外开浏览器。

返回:[回到 README](../README.md)
"""

PY_SRC = '''"""CodeMirror 路径样本 —— 代码类文件走 CodeMirror,选中行为与 DocumentViewer 不同。"""


def marked_function(value):
    """PYMARK 特征函数,便于在预览里定位。"""
    total = 0
    for index in range(value):
        total += index * 2
    return total


class SampleClass:
    def __init__(self, name):
        self.name = name

    def describe(self):
        return f"sample:{self.name}"


if __name__ == "__main__":
    print(marked_function(10))
'''

JSON_SRC = """{
  "marker": "JSONMARK",
  "nested": { "list": [1, 2, 3], "flag": true },
  "chinese": "中文键值也要正常显示"
}
"""

TOML_SRC = """# TOMLMARK
[section]
name = "sample"
count = 42

[section.nested]
enabled = true
"""

TXT_SRC = ("纯文本样本 TXTMARK\n"
           "第二行,含中文与 ASCII 混排 mixed 123\n"
           "第三行\n")


def w(rel, content):
    path = os.path.join(ROOT, rel)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def make_docx():
    from docx import Document
    doc = Document()
    doc.add_heading("DOCX 样本标题 DOCXMARK", level=1)
    doc.add_paragraph("第一段正文,内置 LibreOffice 转换后应能看到这句。")
    doc.add_heading("小节标题", level=2)
    for i in range(1, 4):
        doc.add_paragraph("列表项 %d" % i, style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "表头A"
    table.cell(0, 1).text = "表头B"
    table.cell(1, 0).text = "值1"
    table.cell(1, 1).text = "值2"
    path = os.path.join(ROOT, "docs", "sample.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


def make_xlsx():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "数据表"
    ws.append(["名称", "数量", "备注"])
    for i in range(1, 11):
        ws.append(["条目%d" % i, i * 10, "XLSXMARK" if i == 1 else ""])
    ws2 = wb.create_sheet("第二表")
    ws2.append(["SHEET2MARK"])
    path = os.path.join(ROOT, "docs", "sample.xlsx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    wb.save(path)
    return path


def make_png():
    from PIL import Image, ImageDraw
    img = Image.new("RGB", (640, 360), (36, 62, 99))
    d = ImageDraw.Draw(img)
    # 画三个高对比色块 —— 图片预览是否真渲染可用像素采样判定,不靠「看着有图」
    d.rectangle([40, 40, 200, 200], fill=(220, 80, 60))
    d.rectangle([240, 40, 400, 200], fill=(90, 200, 120))
    d.rectangle([440, 40, 600, 200], fill=(240, 200, 70))
    d.text((40, 260), "PNGMARK 640x360", fill=(255, 255, 255))
    path = os.path.join(ROOT, "images", "sample.png")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    img.save(path)
    return path


def make_pdf(docx_path):
    """用应用内置的 LibreOffice 转 PDF —— 与产品实际使用的是同一套引擎。"""
    if not APP or not os.path.exists(APP):
        print("  ! 找不到可用的 LibreOffice(打包产物内 / 系统安装均无),跳过 PDF 生成")
        return None
    print("  · soffice:%s" % APP)
    out_dir = os.path.join(ROOT, "docs")
    r = subprocess.run([APP, "--headless", "--norestore", "--convert-to", "pdf",
                        "--outdir", out_dir, docx_path],
                       capture_output=True, text=True, timeout=180)
    path = os.path.join(out_dir, "sample.pdf")
    if not os.path.exists(path):
        print("  ! PDF 生成失败:", (r.stderr or r.stdout)[:200])
        return None
    return path


def make_big():
    """大文件预览守卫(#46)的取材。行内带行号,便于确认是截断还是全量。"""
    path = os.path.join(ROOT, "big", "large.txt")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        for i in range(400000):
            f.write("第 %07d 行 —— 大文件守卫样本 BIGMARK 填充填充填充填充\n" % i)
    return path


def git(*args):
    return subprocess.run(["git", "-C", ROOT] + list(args), capture_output=True, text=True)


def main():
    os.makedirs(ROOT, exist_ok=True)
    print("生成测试项目:%s" % ROOT)

    made = [w("README.md", MD_MAIN), w("notes/sub.md", MD_SUB),
            w("code/sample.py", PY_SRC), w("code/data.json", JSON_SRC),
            w("code/config.toml", TOML_SRC), w("code/plain.txt", TXT_SRC)]
    # FORK 2026-08-14 [feat: upstream-sync-2026-08]:中文名 + 名字带空格的样本。
    # Windows 路径处理最容易翻车的就是这两类(编码 + 引号/转义),而原 fixture 全是纯 ASCII,
    # 于是 `win_p0_paths.py` 的「中文名条目」「Ctrl+K 搜中文」两条只能记 SKIP ——
    # 测试项目自己不带这类文件,就只能借 user 的真实项目来验,既不稳定也不该。
    made.append(w("docs/中文文件名 带空格.md",
                  "# 中文样本 CJKMARK\n\n用于验证 Windows 下中文 + 空格路径的文件树 / 预览 / 搜索。\n"))
    made.append(w("docs/嵌套目录/更深一层.md", "# 更深一层 DEEPMARK\n\n验证多级中文目录的路径拼接。\n"))
    docx = make_docx(); made.append(docx)
    made.append(make_xlsx())
    made.append(make_png())
    pdf = make_pdf(docx)
    if pdf:
        made.append(pdf)
    made.append(make_big())

    # git 仓 + 故意留改动:第 2 组 #15 的「N 更改」tab 没有 diff 就不渲染
    if not os.path.isdir(os.path.join(ROOT, ".git")):
        git("init", "-q")
        git("config", "user.email", "uitest@local")
        git("config", "user.name", "uitest")
    w(".gitignore", "big/\n")
    git("add", "-A")
    git("commit", "-q", "-m", "fixtures baseline")
    w("README.md", MD_MAIN + "\n<!-- 未提交改动:让「N 更改」tab 有内容 -->\n")
    w("code/untracked.py", "# 未跟踪文件 UNTRACKEDMARK\nvalue = 1\n")

    st = git("status", "--porcelain").stdout.strip().splitlines()
    print("\n共 %d 个样本文件,git 未提交改动 %d 条" % (len(made), len(st)))
    for p in made:
        print("  %8.1f KB  %s" % (os.path.getsize(p) / 1024, p.replace(ROOT + "/", "")))
    print("\n改动清单:", ", ".join(st))


if __name__ == "__main__":
    main()
